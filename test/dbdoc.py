#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2020/12/29 下午5:19
# @Author  : kiwanter
# @Email   : kiwanter@163.com
# @File    : dbdoc.py
# @software: pycharm

'''
@file function:
'''

import pymysql.cursors
from docx import Document
import config

doc_path = 'dbdoc.docx'
doc = Document()
doc.styles['Normal'].font.name = u'宋体'


def to_doc_table(rows, table):
    """
    (('id', 'bigint(20)', None, 'NO', 'PRI', None, 'auto_increment', 'select,insert,update,references', '递增主键'),...)
    Field, Type, Null, Key, Default, Extra, Privileges, Comment
    :param rows: 列表
    :param table: 表名
    """

    # 标题
    doc.add_heading(table + '表', 3)
    t = doc.add_table(rows=1, cols=6, style='Table Grid')
    head_cells = t.rows[0].cells
    for cell in head_cells:
        # font = cell.paragraphs[0].style.font
        # font.bold = True
        p = cell.paragraphs[0]
        run = p.add_run('dcolor')
        run.font.bold = True

    header = ['字段名称', '字段名', '数据类型', '是否可空', '默认值', '备注']
    for i in range(len(header)):
        head_cells[i].text = header[i]

    # 内容
    for r in rows:
        row = t.add_row().cells
        cols = [r[8], r[0], r[1], r[3], r[4], r[8]]
        for i in range(len(cols)):
            # row[i].paragraphs[0].style.font.bold = False
            row[i].text = cols[i] if cols[i] else ''

    # doc.add_page_break()
    doc.save(doc_path)


def generate():
    con = pymysql.connect(
        host='127.0.0.1',
        user='root',
        password='123456',
        db='exam'
    )
    try:
        with con.cursor() as cursor:
            sql = 'show tables'
            cursor.execute(sql)
            tables = cursor.fetchall()
            print(tables)
            for t in tables:
                # desc table 不能展示注释
                sql = 'show full columns from ' + t[0]
                cursor.execute(sql)
                rows = cursor.fetchall()
                print(rows)
                to_doc_table(rows, t[0])
    finally:
        con.close()


if __name__ == '__main__':
    """
    pip3 install python-docx
    文档参考：https://buildmedia.readthedocs.org/media/pdf/python-docx/latest/python-docx.pdf
    pip3 install PyMySQL
    https://github.com/PyMySQL/PyMySQL
    """
    generate()